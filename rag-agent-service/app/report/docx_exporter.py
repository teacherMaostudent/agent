"""Markdown → Word(.docx) 转换器。

把审查报告 / 生成初稿的 Markdown 转成排好版的 Word 文档，让用户打开就是
标题、加粗、表格、列表等原生格式，看不到 #、**、---、| 这类 Markdown 符号。

只覆盖本项目实际产出的语法(标题 #/##/###、加粗 **、无序列表 -、有序列表 1.、
表格 |、引用 >、分隔线 ---)，不追求完整 Markdown 规范，够用且好维护。
用已安装的 python-docx，无需额外依赖。
"""

import io
import re

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

# 合规免责声明：所有导出文件开头必须加,明确 AI 结果仅供参考、不构成最终合规结论。
DISCLAIMER = (
    "以上分析结果基于AI对输入文件的自动化比对，仅供参考与预警，"
    "不构成最终合规结论。具体判定请以企业质量管理部门/QA意见为准。"
)


# 行首内联加粗：**xxx** → 加粗 run。返回 (纯文本片段, 是否加粗) 列表。
_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _emit_run(paragraph, text: str, highlight: list[str] | None, bold: bool = False) -> None:
    """写入一段文字。若命中 highlight 里的词，则该词加粗 + 黄底，其余正常。

    highlight 里的词按长度降序匹配，避免短词先切断长词。
    """
    if not text:
        return
    if not highlight:
        run = paragraph.add_run(text)
        run.bold = bold
        return
    # 用正则一次找出所有高亮词的位置，逐段写入。
    pattern = "|".join(re.escape(w) for w in sorted(highlight, key=len, reverse=True) if w)
    if not pattern:
        run = paragraph.add_run(text)
        run.bold = bold
        return
    pos = 0
    for m in re.finditer(pattern, text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()]).bold = bold
        hit = paragraph.add_run(m.group(0))
        hit.bold = True
        hit.font.highlight_color = WD_COLOR_INDEX.YELLOW
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:]).bold = bold


def _add_runs(paragraph, text: str, highlight: list[str] | None = None) -> None:
    """把一行文字按 **加粗** 切成若干 run 写入段落，去掉 ** 符号。

    highlight：需在正文中加粗+黄底标出的词（如待人工确认的模糊词、混用术语）。
    """
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            _emit_run(paragraph, text[pos : m.start()], highlight)
        _emit_run(paragraph, m.group(1), highlight, bold=True)
        pos = m.end()
    if pos < len(text):
        _emit_run(paragraph, text[pos:], highlight)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_separator_row(line: str) -> bool:
    """表格分隔行，如 |---|---| 或 | :--- | ---: |，不作为数据行渲染。"""
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def markdown_to_docx_bytes(
    markdown: str, title: str = "", highlight: list[str] | None = None
) -> bytes:
    """把 Markdown 文本转成 .docx 的字节流，可直接作为下载响应体。

    highlight：正文中需要加粗+黄底标出的词（待人工确认的问题点）。
    """
    highlight = [w for w in (highlight or []) if w and w.strip()]
    doc = Document()
    # 文档最前面加合规免责声明(灰色小字),所有导出的 Word 都带。
    disc = doc.add_paragraph()
    disc_run = disc.add_run(DISCLAIMER)
    disc_run.italic = True
    disc_run.font.size = Pt(9)
    disc_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 分隔线 --- / *** → Word 分页级分隔，用一条空段落+浅色横线替代（简化为空段落）
        if re.fullmatch(r"[-*_]{3,}", stripped):
            doc.add_paragraph()
            i += 1
            continue

        # 标题 #/##/###...
        heading = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(heading.group(2).strip(), level=level)
            i += 1
            continue

        # 表格：连续的 | 行聚成一张表
        if _is_table_row(line):
            block: list[str] = []
            while i < n and _is_table_row(lines[i]):
                block.append(lines[i])
                i += 1
            _render_table(doc, block, highlight)
            continue

        # 引用 >
        if stripped.startswith(">"):
            quote = stripped.lstrip(">").strip()
            p = (
                doc.add_paragraph(style="Intense Quote")
                if _has_style(doc, "Intense Quote")
                else doc.add_paragraph()
            )
            _add_runs(p, quote, highlight)
            i += 1
            continue

        # 有序列表 1. 2. 3.
        ordered = re.match(r"\d+\.\s+(.*)", stripped)
        if ordered:
            p = (
                doc.add_paragraph(style="List Number")
                if _has_style(doc, "List Number")
                else doc.add_paragraph()
            )
            _add_runs(p, ordered.group(1), highlight)
            i += 1
            continue

        # 无序列表 - / * / +（带缩进的也归为列表项）
        bullet = re.match(r"[-*+]\s+(.*)", stripped)
        if bullet:
            p = (
                doc.add_paragraph(style="List Bullet")
                if _has_style(doc, "List Bullet")
                else doc.add_paragraph()
            )
            _add_runs(p, bullet.group(1), highlight)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_runs(p, stripped, highlight)
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_table(doc: Document, block: list[str], highlight: list[str] | None = None) -> None:
    """把 | 分隔的行块渲染成 Word 表格，跳过 |---| 分隔行，首行作表头加粗。"""
    rows = [r for r in block if not _is_separator_row(r)]
    if not rows:
        return
    matrix = [_split_row(r) for r in rows]
    cols = max(len(r) for r in matrix)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid" if _has_style(doc, "Table Grid") else None
    for r_idx, cells in enumerate(matrix):
        cells = cells + [""] * (cols - len(cells))  # 补齐短行
        row = table.add_row()
        for c_idx, cell_text in enumerate(cells):
            cell = row.cells[c_idx]
            cell.text = ""  # 清掉默认空段落文本
            _add_runs(cell.paragraphs[0], cell_text, highlight)
            if r_idx == 0:  # 表头加粗
                for run in cell.paragraphs[0].runs:
                    run.bold = True


def _has_style(doc: Document, name: str) -> bool:
    try:
        _ = doc.styles[name]
        return True
    except KeyError:
        return False
